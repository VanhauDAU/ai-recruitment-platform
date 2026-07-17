from copy import deepcopy
from hashlib import sha256
from datetime import timedelta
from io import BytesIO
import shutil
import tempfile
from unittest.mock import patch

from django.contrib.auth import get_user_model
from django.core.files.base import ContentFile
from django.core.files.uploadedfile import SimpleUploadedFile
from django.core.files.storage import default_storage
from django.test import override_settings
from django.urls import reverse
from django.utils import timezone
from rest_framework.test import APITestCase
from PIL import Image
from pypdf import PdfReader

from apps.cv_templates.models import (
    CvColor,
    CvContentBlueprint,
    CvSampleContent,
    CvTemplate,
    CvTemplateColorLink,
    CvTemplateVersion,
)
from apps.jobs.models import JobCategory, JobCategoryLocalization

from .composition import compose_cv_document
from .models import CvAccessLog, CvAsset, CvDraft, CvExport, CvImportJob, CvSharedLink, CvVersion, UserCv
from .pdf_renderer import build_cv_pdf_html, render_cv_version_pdf
from .schemas import empty_content, empty_layout, empty_style
from .services.versions import sync_legacy_builder_draft
from .tasks import ImportProcessingError, generate_cv_thumbnail, process_cv_import_job, render_cv_export_job


TEST_EXPORT_MEDIA_ROOT = tempfile.mkdtemp()


@override_settings(MEDIA_ROOT=TEST_EXPORT_MEDIA_ROOT, ALLOWED_HOSTS=['testserver'])
class CvV2ApiTests(APITestCase):
    @classmethod
    def tearDownClass(cls):
        super().tearDownClass()
        shutil.rmtree(TEST_EXPORT_MEDIA_ROOT, ignore_errors=True)

    def setUp(self):
        self.candidate = get_user_model().objects.create_user(
            email='v2-candidate@example.com', password='password', role='candidate', email_verified=True,
        )
        self.other_candidate = get_user_model().objects.create_user(
            email='v2-other@example.com', password='password', role='candidate', email_verified=True,
        )
        self.template = CvTemplate.objects.create(
            name='V2 Template', lifecycle_status=CvTemplate.LifecycleStatus.PUBLISHED,
        )
        self.template_version = CvTemplateVersion.objects.create(
            template=self.template,
            version_number=1,
            version_status=CvTemplateVersion.VersionStatus.PUBLISHED,
            renderer_key='classic_single_column_v1',
            renderer_version='1',
            default_layout_json=empty_layout(),
            default_style_json=empty_style(),
            content_contract={'schema': 'canonical_cv_content_v1', 'schema_version': 1},
        )
        self.template.current_published_version = self.template_version
        self.template.save(update_fields=['current_published_version'])
        self.blue = CvColor.objects.create(name='Blue', slug='blue', hex_code='#2255AA')
        CvTemplateColorLink.objects.create(template=self.template, color=self.blue, is_default=True)
        self.client.force_authenticate(self.candidate)

    def create_cv(self):
        response = self.client.post(
            reverse('cv-v2-list-create'),
            {'title': 'CV V2', 'template_public_id': self.template.public_id, 'language': 'vi-VN'},
            format='json',
        )
        self.assertEqual(response.status_code, 201, response.data)
        return UserCv.objects.get(public_id=response.data['public_id'])

    def draft_url(self, cv):
        return reverse('cv-v2-draft', kwargs={'public_id': cv.public_id})

    def owner_view_url(self, cv):
        return reverse('cv-v2-owner-view', kwargs={'public_id': cv.public_id})

    def shared_links_url(self, cv):
        return reverse('cv-v2-shared-link-list-create', kwargs={'public_id': cv.public_id})

    def exports_url(self, cv):
        return reverse('cv-v2-export-list-create', kwargs={'public_id': cv.public_id})

    def export_url(self, cv, export, action=None):
        name = {
            None: 'cv-v2-export-detail',
            'retry': 'cv-v2-export-retry',
            'download': 'cv-v2-export-download',
        }[action]
        return reverse(name, kwargs={'public_id': cv.public_id, 'export_public_id': export.public_id})

    def create_alternate_template(self):
        template = CvTemplate.objects.create(
            name='Two column V2 Template', lifecycle_status=CvTemplate.LifecycleStatus.PUBLISHED,
        )
        layout = empty_layout()
        layout['regions'] = [
            {'id': 'main', 'width_percent': 68, 'section_instance_ids': []},
            {'id': 'sidebar', 'width_percent': 32, 'section_instance_ids': []},
        ]
        style = {**empty_style(), 'theme_color': '#2255AA', 'font_family': 'Inter'}
        version = CvTemplateVersion.objects.create(
            template=template,
            version_number=1,
            version_status=CvTemplateVersion.VersionStatus.PUBLISHED,
            renderer_key='classic_two_column_v1',
            renderer_version='1',
            default_layout_json=layout,
            default_style_json=style,
            capabilities={
                'layout': {
                    'section_drag': True,
                    'cross_region_drag': True,
                    'item_drag': True,
                    'column_resize': {'enabled': True, 'min_percent': 25, 'max_percent': 75},
                },
            },
        )
        template.current_published_version = version
        template.save(update_fields=['current_published_version'])
        return template, version

    def test_canonical_composer_applies_presentation_without_mutating_content(self):
        content = empty_content('vi-VN')
        original = deepcopy(content)

        document = compose_cv_document(
            template=self.template,
            content_json=content,
            theme_color='#2255AA',
        )

        self.assertEqual(content, original)
        self.assertIsNot(document['content_json'], content)
        self.assertEqual(document['style_json']['theme_color'], '#2255AA')
        self.assertEqual(document['schema_version'], self.template_version.schema_version)

    def test_create_autosave_save_and_publish_follow_the_lifecycle_contract(self):
        cv = self.create_cv()
        self.assertEqual(CvVersion.objects.filter(cv=cv).count(), 1)

        detail = self.client.get(reverse('cv-v2-detail', kwargs={'public_id': cv.public_id}))
        self.assertEqual(detail.status_code, 200)
        self.assertEqual(detail.data['template_renderer_key'], 'classic_single_column_v1')
        self.assertEqual(detail.data['template_renderer_version'], '1')

        owner_view = self.client.get(self.owner_view_url(cv))
        self.assertEqual(owner_view.status_code, 200)
        self.assertIn('is_default', owner_view.data['cv'])

        draft_response = self.client.get(self.draft_url(cv))
        self.assertEqual(draft_response.status_code, 200)
        self.assertEqual(draft_response['ETag'], '"lock-version-0"')
        payload = {
            key: draft_response.data[key]
            for key in ('schema_version', 'content_json', 'layout_json', 'style_json')
        }
        payload['content_json']['personal_info']['full_name'] = 'Candidate V2'

        autosave = self.client.put(
            self.draft_url(cv), payload, format='json', HTTP_IF_MATCH='"lock-version-0"',
        )
        self.assertEqual(autosave.status_code, 200, autosave.data)
        self.assertEqual(autosave.data['lock_version'], 1)
        self.assertEqual(CvVersion.objects.filter(cv=cv).count(), 1)
        cv.refresh_from_db()
        self.assertEqual(cv.cv_data['personal_info']['full_name'], 'Candidate V2')

        stale = self.client.put(
            self.draft_url(cv), payload, format='json', HTTP_IF_MATCH='"lock-version-0"',
        )
        self.assertEqual(stale.status_code, 409)
        self.assertEqual(stale.data['current_lock_version'], 1)

        saved = self.client.post(
            reverse('cv-v2-save-version', kwargs={'public_id': cv.public_id}),
            format='json', HTTP_IF_MATCH='"lock-version-1"',
        )
        self.assertEqual(saved.status_code, 201, saved.data)
        self.assertEqual(saved.data['version_kind'], CvVersion.VersionKind.MANUAL_SAVE)
        self.assertEqual(CvVersion.objects.filter(cv=cv).count(), 2)

        published = self.client.post(
            reverse('cv-v2-publish', kwargs={'public_id': cv.public_id}),
            format='json', HTTP_IF_MATCH='"lock-version-1"',
        )
        self.assertEqual(published.status_code, 201, published.data)
        self.assertEqual(published.data['version_kind'], CvVersion.VersionKind.PUBLISHED)
        cv.refresh_from_db()
        self.assertEqual(cv.lifecycle_status, UserCv.LifecycleStatus.PUBLISHED)
        self.assertEqual(cv.published_version.public_id, published.data['public_id'])

    def test_latest_recoverable_draft_uses_hash_and_clears_after_manual_save(self):
        cv = self.create_cv()
        url = reverse('cv-v2-latest-recoverable-draft')
        self.assertEqual(self.client.get(url).status_code, 204)

        draft_response = self.client.get(self.draft_url(cv))
        payload = {
            key: draft_response.data[key]
            for key in ('schema_version', 'content_json', 'layout_json', 'style_json')
        }
        payload['content_json']['personal_info']['full_name'] = 'Recover me'
        autosave = self.client.put(
            self.draft_url(cv), payload, format='json', HTTP_IF_MATCH='"lock-version-0"',
        )
        self.assertEqual(autosave.status_code, 200, autosave.data)

        recoverable = self.client.get(url)
        self.assertEqual(recoverable.status_code, 200, recoverable.data)
        self.assertEqual(recoverable.data['cv']['public_id'], cv.public_id)
        self.assertEqual(recoverable.data['draft']['content_json']['personal_info']['full_name'], 'Recover me')

        saved = self.client.post(
            reverse('cv-v2-save-version', kwargs={'public_id': cv.public_id}),
            format='json',
            HTTP_IF_MATCH='"lock-version-1"',
        )
        self.assertEqual(saved.status_code, 201, saved.data)
        self.assertEqual(self.client.get(url).status_code, 204)

    def test_latest_recoverable_draft_returns_only_most_recent_dirty_cv(self):
        first = self.create_cv()
        second = self.create_cv()
        for cv, name in ((first, 'Older draft'), (second, 'Newest draft')):
            draft_response = self.client.get(self.draft_url(cv))
            payload = {
                key: draft_response.data[key]
                for key in ('schema_version', 'content_json', 'layout_json', 'style_json')
            }
            payload['content_json']['personal_info']['full_name'] = name
            response = self.client.put(
                self.draft_url(cv), payload, format='json', HTTP_IF_MATCH='"lock-version-0"',
            )
            self.assertEqual(response.status_code, 200, response.data)

        recoverable = self.client.get(reverse('cv-v2-latest-recoverable-draft'))
        self.assertEqual(recoverable.status_code, 200, recoverable.data)
        self.assertEqual(recoverable.data['cv']['public_id'], second.public_id)

    def test_create_can_copy_owned_draft_into_a_new_template(self):
        source = self.create_cv()
        draft_response = self.client.get(self.draft_url(source))
        payload = {
            key: draft_response.data[key]
            for key in ('schema_version', 'content_json', 'layout_json', 'style_json')
        }
        payload['content_json']['personal_info']['full_name'] = 'Preserved source identity'
        self.client.put(
            self.draft_url(source), payload, format='json', HTTP_IF_MATCH='"lock-version-0"',
        )
        alternate, _ = self.create_alternate_template()

        response = self.client.post(reverse('cv-v2-list-create'), {
            'title': 'Copied CV',
            'template_public_id': alternate.public_id,
            'source_cv_public_id': source.public_id,
            'language': 'en-US',
        }, format='json')

        self.assertEqual(response.status_code, 201, response.data)
        copied = UserCv.objects.get(public_id=response.data['public_id'])
        self.assertNotEqual(copied.pk, source.pk)
        self.assertEqual(copied.language, 'vi-VN')
        self.assertEqual(copied.draft.content_json['personal_info']['full_name'], 'Preserved source identity')
        self.assertEqual(copied.template_id, alternate.pk)

    def test_template_preview_projects_draft_without_mutating_cv(self):
        cv = self.create_cv()
        original_template_id = cv.template_id
        alternate, version = self.create_alternate_template()

        response = self.client.get(
            reverse('cv-v2-template-preview', kwargs={'public_id': cv.public_id}),
            {'template_public_id': alternate.public_id},
        )

        self.assertEqual(response.status_code, 200, response.data)
        self.assertEqual(response.data['renderer']['key'], version.renderer_key)
        cv.refresh_from_db()
        self.assertEqual(cv.template_id, original_template_id)

    def test_create_applies_only_a_color_available_for_the_template(self):
        response = self.client.post(
            reverse('cv-v2-list-create'),
            {
                'title': 'Blue CV',
                'template_public_id': self.template.public_id,
                'language': 'vi-VN',
                'theme_color': '#2255aa',
            },
            format='json',
        )

        self.assertEqual(response.status_code, 201, response.data)
        cv = UserCv.objects.get(public_id=response.data['public_id'])
        self.assertEqual(cv.draft.style_json['theme_color'], '#2255AA')
        self.assertEqual(cv.latest_version.style_json['theme_color'], '#2255AA')

        invalid = self.client.post(
            reverse('cv-v2-list-create'),
            {
                'title': 'Invalid color CV',
                'template_public_id': self.template.public_id,
                'language': 'vi-VN',
                'theme_color': '#FF0000',
            },
            format='json',
        )
        self.assertEqual(invalid.status_code, 400)
        self.assertIn('theme_color', invalid.data)

    def test_import_hard_delete_and_metadata_use_v2_without_legacy_storage_urls(self):
        uploaded = SimpleUploadedFile('source.pdf', b'%PDF-1.4 test', content_type='application/pdf')
        imported = self.client.post(
            reverse('cv-v2-import'),
            {'file': uploaded, 'title': 'Imported CV'},
            format='multipart',
        )

        self.assertEqual(imported.status_code, 201, imported.data)
        self.assertEqual(imported.data['source'], UserCv.Source.IMPORTED)
        self.assertEqual(imported.data['cv_type'], UserCv.CvType.UPLOADED)
        self.assertEqual(imported.data['file_name'], 'source.pdf')
        self.assertNotIn('file_url', imported.data)
        cv = UserCv.objects.get(public_id=imported.data['public_id'])
        self.assertEqual(cv.latest_version.version_kind, CvVersion.VersionKind.IMPORTED)

        metadata = self.client.patch(
            reverse('cv-v2-detail', kwargs={'public_id': cv.public_id}),
            {'title': 'Renamed CV', 'is_default': True},
            format='json',
        )
        self.assertEqual(metadata.status_code, 200, metadata.data)
        self.assertEqual(metadata.data['title'], 'Renamed CV')
        self.assertTrue(metadata.data['is_default'])

        deleted = self.client.delete(reverse('cv-v2-detail', kwargs={'public_id': cv.public_id}))
        self.assertEqual(deleted.status_code, 204)
        self.assertFalse(UserCv.objects.filter(pk=cv.pk).exists())
        self.assertFalse(CvVersion.objects.filter(cv_id=cv.pk).exists())
        self.assertEqual(
            self.client.get(reverse('cv-v2-detail', kwargs={'public_id': cv.public_id})).status_code,
            404,
        )

    def _docx_upload(self, name='source.docx'):
        from docx import Document

        document = Document()
        document.add_paragraph('Nguyen Van Candidate')
        document.add_paragraph('candidate@example.com | +84 912 345 678')
        document.add_paragraph('Software engineer with five years of product experience.')
        output = BytesIO()
        document.save(output)
        return SimpleUploadedFile(
            name, output.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )

    @patch('apps.cvs.tasks.structure_cv_text')
    def test_template_import_is_idempotent_and_materializes_editable_canonical_draft(self, structure_cv_text):
        payload = {
            'file': self._docx_upload(),
            'title': 'Parsed CV',
            'template_public_id': self.template.public_id,
            'language': 'vi-VN',
            'theme_color': '#2255AA',
        }
        with patch('apps.cvs.tasks.process_cv_import_job.delay'):
            first = self.client.post(
                reverse('cv-v2-import'), payload, format='multipart',
                HTTP_IDEMPOTENCY_KEY='same-import',
            )
            second = self.client.post(
                reverse('cv-v2-import'), {**payload, 'file': self._docx_upload()}, format='multipart',
                HTTP_IDEMPOTENCY_KEY='same-import',
            )

        self.assertEqual(first.status_code, 202, first.data)
        self.assertEqual(second.status_code, 200, second.data)
        self.assertEqual(first.data['public_id'], second.data['public_id'])
        self.assertEqual(CvImportJob.objects.count(), 1)

        cv = UserCv.objects.get(public_id=first.data['public_id'])
        parsed_content = empty_content('vi-VN')
        parsed_content['personal_info']['full_name'] = 'Nguyen Van Candidate'
        structure_cv_text.return_value = parsed_content
        process_cv_import_job(cv.import_job.pk)
        cv.refresh_from_db()
        cv.import_job.refresh_from_db()
        self.assertEqual(cv.processing_status, UserCv.ProcessingStatus.ANALYZED)
        self.assertEqual(cv.import_job.status, CvImportJob.Status.COMPLETED)
        self.assertEqual(cv.latest_version.version_kind, CvVersion.VersionKind.IMPORTED)
        self.assertEqual(cv.draft.content_json['personal_info']['full_name'], 'Nguyen Van Candidate')
        self.assertEqual(cv.draft.style_json['theme_color'], '#2255AA')
        self.assertEqual(cv.raw_text, '')

    def test_import_rejects_spoofed_file_signature_before_persisting(self):
        upload = SimpleUploadedFile('spoofed.pdf', b'not a real pdf', content_type='application/pdf')
        response = self.client.post(
            reverse('cv-v2-import'),
            {'file': upload, 'template_public_id': self.template.public_id, 'language': 'vi-VN'},
            format='multipart',
        )

        self.assertEqual(response.status_code, 400)
        self.assertEqual(CvImportJob.objects.count(), 0)

    def test_failed_import_exposes_safe_code_and_can_retry(self):
        with patch('apps.cvs.tasks.process_cv_import_job.delay'):
            queued = self.client.post(
                reverse('cv-v2-import'),
                {
                    'file': self._docx_upload(),
                    'template_public_id': self.template.public_id,
                    'language': 'vi-VN',
                },
                format='multipart',
            )
        cv = UserCv.objects.get(public_id=queued.data['public_id'])
        with patch('apps.cvs.tasks._extract_import_text', side_effect=ImportProcessingError('scanned_pdf_ocr_unavailable')):
            process_cv_import_job(cv.import_job.pk)

        cv.refresh_from_db()
        cv.import_job.refresh_from_db()
        self.assertEqual(cv.processing_status, UserCv.ProcessingStatus.FAILED)
        self.assertEqual(cv.import_job.failure_code, 'scanned_pdf_ocr_unavailable')
        self.assertNotIn('Nguyen Van Candidate', cv.error_message)

        with patch('apps.cvs.tasks.process_cv_import_job.delay') as enqueue:
            with self.captureOnCommitCallbacks(execute=True):
                retried = self.client.post(
                    reverse('cv-v2-import-retry', kwargs={'public_id': cv.public_id}),
                    {}, format='json',
                )
        self.assertEqual(retried.status_code, 202, retried.data)
        cv.import_job.refresh_from_db()
        self.assertEqual(cv.import_job.status, CvImportJob.Status.QUEUED)
        enqueue.assert_called_once_with(cv.import_job.pk)

    def test_setting_another_default_cv_clears_the_previous_active_default(self):
        first_cv = self.create_cv()
        second_cv = self.create_cv()
        first = self.client.patch(
            reverse('cv-v2-detail', kwargs={'public_id': first_cv.public_id}),
            {'is_default': True},
            format='json',
        )
        second = self.client.patch(
            reverse('cv-v2-detail', kwargs={'public_id': second_cv.public_id}),
            {'is_default': True},
            format='json',
        )

        self.assertEqual(first.status_code, 200, first.data)
        self.assertEqual(second.status_code, 200, second.data)
        first_cv.refresh_from_db()
        second_cv.refresh_from_db()
        self.assertFalse(first_cv.is_default)
        self.assertTrue(second_cv.is_default)

    def test_duplicate_creates_an_independent_builder_draft_from_the_latest_version(self):
        source_cv = self.create_cv()
        source_draft = source_cv.draft
        source_draft.content_json['personal_info']['full_name'] = 'Original candidate'
        source_draft.save(update_fields=['content_json'])
        saved = self.client.post(
            reverse('cv-v2-save-version', kwargs={'public_id': source_cv.public_id}),
            format='json',
            HTTP_IF_MATCH='"lock-version-0"',
        )
        self.assertEqual(saved.status_code, 201, saved.data)
        source_cv.refresh_from_db()

        response = self.client.post(
            reverse('cv-v2-duplicate', kwargs={'public_id': source_cv.public_id}),
            {'title': 'CV duplicated'},
            format='json',
        )

        self.assertEqual(response.status_code, 201, response.data)
        duplicate = UserCv.objects.get(public_id=response.data['public_id'])
        self.assertNotEqual(duplicate.pk, source_cv.pk)
        self.assertEqual(duplicate.title, 'CV duplicated')
        self.assertEqual(duplicate.lifecycle_status, UserCv.LifecycleStatus.DRAFT)
        self.assertEqual(duplicate.latest_version.version_kind, CvVersion.VersionKind.INITIAL)
        self.assertEqual(duplicate.current_template_version_id, source_cv.current_template_version_id)
        self.assertEqual(duplicate.draft.content_json['personal_info']['full_name'], 'Original candidate')

        duplicate_draft = duplicate.draft
        duplicate_draft.content_json['personal_info']['full_name'] = 'Duplicate candidate'
        duplicate_draft.save(update_fields=['content_json'])
        source_cv.refresh_from_db()
        self.assertEqual(source_cv.draft.content_json['personal_info']['full_name'], 'Original candidate')

    def test_delete_permanently_removes_the_cv_from_owner_library(self):
        cv = self.create_cv()
        self.assertEqual(self.client.delete(reverse('cv-v2-detail', kwargs={'public_id': cv.public_id})).status_code, 204)
        self.assertFalse(UserCv.objects.filter(pk=cv.pk).exists())
        listed = self.client.get(reverse('cv-v2-list-create')).data
        self.assertEqual(listed['results'] if isinstance(listed, dict) else listed, [])

    def test_autosave_rejects_invalid_canonical_documents_before_writing(self):
        cv = self.create_cv()
        draft = self.client.get(self.draft_url(cv)).data
        payload = {
            key: draft[key]
            for key in ('schema_version', 'content_json', 'layout_json', 'style_json')
        }
        payload['style_json']['theme_color'] = 'not-a-color'

        response = self.client.put(
            self.draft_url(cv), payload, format='json', HTTP_IF_MATCH='"lock-version-0"',
        )

        self.assertEqual(response.status_code, 400)
        cv.refresh_from_db()
        self.assertEqual(cv.draft.lock_version, 0)
        self.assertEqual(CvVersion.objects.filter(cv=cv).count(), 1)

    def test_save_and_publish_validate_the_draft_before_creating_a_version(self):
        cv = self.create_cv()
        CvDraft.objects.filter(cv=cv).update(style_json={**empty_style(), 'theme_color': 'invalid'})

        save = self.client.post(
            reverse('cv-v2-save-version', kwargs={'public_id': cv.public_id}),
            format='json', HTTP_IF_MATCH='"lock-version-0"',
        )
        publish = self.client.post(
            reverse('cv-v2-publish', kwargs={'public_id': cv.public_id}),
            format='json', HTTP_IF_MATCH='"lock-version-0"',
        )

        self.assertEqual(save.status_code, 400)
        self.assertEqual(publish.status_code, 400)
        self.assertEqual(CvVersion.objects.filter(cv=cv).count(), 1)

    def test_template_switch_preserves_content_and_changes_only_mutable_presentation(self):
        cv = self.create_cv()
        alternate_template, alternate_version = self.create_alternate_template()
        original_version = cv.latest_version
        original_content = deepcopy(cv.draft.content_json)

        response = self.client.put(
            reverse('cv-v2-template-switch', kwargs={'public_id': cv.public_id}),
            {'template_public_id': alternate_template.public_id, 'client_session_id': 'editor-tab'},
            format='json', HTTP_IF_MATCH='"lock-version-0"',
        )

        self.assertEqual(response.status_code, 200, response.data)
        self.assertEqual(response['ETag'], '"lock-version-1"')
        self.assertEqual(response.data['draft']['content_json'], original_content)
        self.assertEqual(response.data['cv']['template_renderer_key'], 'classic_two_column_v1')
        self.assertEqual(response.data['cv']['template_capabilities']['layout']['column_resize']['min_percent'], 25)
        self.assertEqual(response.data['draft']['layout_json']['regions'][1]['id'], 'sidebar')
        self.assertEqual(response.data['draft']['style_json']['theme_color'], '#2255AA')
        cv.refresh_from_db()
        self.assertEqual(cv.template_id, alternate_template.id)
        self.assertEqual(cv.current_template_version_id, alternate_version.id)
        self.assertEqual(cv.latest_version_id, original_version.id)

        invalid_layout = {
            key: deepcopy(response.data['draft'][key])
            for key in ('schema_version', 'content_json', 'layout_json', 'style_json')
        }
        invalid_layout['layout_json']['regions'][0]['width_percent'] = 90
        invalid_layout['layout_json']['regions'][1]['width_percent'] = 10
        invalid_resize = self.client.put(
            self.draft_url(cv), invalid_layout, format='json', HTTP_IF_MATCH='"lock-version-1"',
        )
        self.assertEqual(invalid_resize.status_code, 400)
        self.assertIn('layout_json.regions', invalid_resize.data)

        stale = self.client.put(
            reverse('cv-v2-template-switch', kwargs={'public_id': cv.public_id}),
            {'template_public_id': self.template.public_id}, format='json', HTTP_IF_MATCH='"lock-version-0"',
        )
        self.assertEqual(stale.status_code, 409)

    def test_layout_item_order_is_validated_and_autosave_keeps_canonical_item_data(self):
        cv = self.create_cv()
        draft = self.client.get(self.draft_url(cv)).data
        payload = {
            key: deepcopy(draft[key])
            for key in ('schema_version', 'content_json', 'layout_json', 'style_json')
        }
        payload['content_json']['sections'] = [{
            'instance_id': 'experience_1',
            'section_key': 'experience',
            'title': 'Kinh nghiệm',
            'enabled': True,
            'items': [
                {'item_id': 'experience_item_1', 'role': 'First', 'company': '', 'start_date': None, 'end_date': None},
                {'item_id': 'experience_item_2', 'role': 'Second', 'company': '', 'start_date': None, 'end_date': None},
            ],
        }]
        payload['layout_json']['regions'][0]['section_instance_ids'] = ['experience_1']
        payload['layout_json']['item_orders'] = {
            'experience_1': ['experience_item_2', 'experience_item_1'],
        }
        original_content = deepcopy(payload['content_json'])

        autosave = self.client.put(
            self.draft_url(cv), payload, format='json', HTTP_IF_MATCH='"lock-version-0"',
        )

        self.assertEqual(autosave.status_code, 200, autosave.data)
        self.assertEqual(autosave.data['content_json'], original_content)
        self.assertEqual(autosave.data['layout_json']['item_orders']['experience_1'], ['experience_item_2', 'experience_item_1'])

        payload['layout_json']['item_orders']['experience_1'] = ['experience_item_1']
        invalid = self.client.put(
            self.draft_url(cv), payload, format='json', HTTP_IF_MATCH='"lock-version-1"',
        )
        self.assertEqual(invalid.status_code, 400)

    def test_owner_view_and_shared_link_are_pinned_to_an_immutable_version_and_audited(self):
        cv = self.create_cv()
        draft = self.client.get(self.draft_url(cv)).data
        save_payload = {
            key: deepcopy(draft[key])
            for key in ('schema_version', 'content_json', 'layout_json', 'style_json')
        }
        save_payload['content_json']['personal_info']['full_name'] = 'Immutable Candidate'
        self.client.put(self.draft_url(cv), save_payload, format='json', HTTP_IF_MATCH='"lock-version-0"')
        saved = self.client.post(
            reverse('cv-v2-save-version', kwargs={'public_id': cv.public_id}),
            format='json', HTTP_IF_MATCH='"lock-version-1"',
        )
        self.assertEqual(saved.status_code, 201, saved.data)

        mutable_payload = deepcopy(save_payload)
        mutable_payload['content_json']['personal_info']['full_name'] = 'Mutable Draft Only'
        self.client.put(self.draft_url(cv), mutable_payload, format='json', HTTP_IF_MATCH='"lock-version-1"')

        owner_view = self.client.get(
            self.owner_view_url(cv), REMOTE_ADDR='203.0.113.9', HTTP_USER_AGENT='owner-browser',
        )
        self.assertEqual(owner_view.status_code, 200, owner_view.data)
        self.assertEqual(owner_view.data['version']['public_id'], saved.data['public_id'])
        self.assertEqual(owner_view.data['version']['content_json']['personal_info']['full_name'], 'Immutable Candidate')
        self.assertNotIn('draft', owner_view.data)

        created = self.client.post(
            self.shared_links_url(cv),
            {'version_public_id': saved.data['public_id']},
            format='json',
        )
        self.assertEqual(created.status_code, 201, created.data)
        raw_token = created.data['token']
        self.assertGreaterEqual(len(raw_token), 43)  # token_urlsafe(32): at least 256 bits of entropy
        link = CvSharedLink.objects.get(public_id=created.data['link']['public_id'])
        self.assertEqual(link.version.public_id, saved.data['public_id'])
        self.assertEqual(link.token_hash, sha256(raw_token.encode('utf-8')).hexdigest())
        self.assertNotEqual(link.token_hash, raw_token)
        self.assertNotIn('token', created.data['link'])

        shared = self.client.get(
            reverse('cv-v2-shared-link-public', kwargs={'token': raw_token}),
            REMOTE_ADDR='198.51.100.4', HTTP_USER_AGENT='shared-browser',
        )
        self.assertEqual(shared.status_code, 200, shared.data)
        self.assertEqual(shared.data['version']['public_id'], saved.data['public_id'])
        self.assertEqual(shared.data['version']['content_json']['personal_info']['full_name'], 'Immutable Candidate')
        self.assertNotIn('plain_text', shared.data['version'])

        logs = CvAccessLog.objects.filter(cv=cv).order_by('accessed_at')
        self.assertEqual(logs.count(), 2)
        owner_log, shared_log = logs
        self.assertEqual(owner_log.access_channel, CvAccessLog.AccessChannel.OWNER_VIEW)
        self.assertEqual(owner_log.actor_user_id, self.candidate.id)
        self.assertEqual(shared_log.access_channel, CvAccessLog.AccessChannel.SHARED_LINK)
        self.assertEqual(shared_log.shared_link_id, link.id)
        self.assertEqual(len(shared_log.ip_hash), 64)
        self.assertEqual(len(shared_log.user_agent_hash), 64)
        self.assertNotEqual(shared_log.ip_hash, '198.51.100.4')
        self.assertNotEqual(shared_log.user_agent_hash, 'shared-browser')
        self.assertNotIn('content_json', {field.name for field in CvAccessLog._meta.fields})

    def test_shared_link_expiry_revoke_and_invalid_tokens_return_404(self):
        cv = self.create_cv()
        created = self.client.post(self.shared_links_url(cv), format='json')
        self.assertEqual(created.status_code, 201, created.data)
        raw_token = created.data['token']
        public_url = reverse('cv-v2-shared-link-public', kwargs={'token': raw_token})
        self.assertEqual(self.client.get(public_url).status_code, 200)

        revoke = self.client.delete(
            reverse('cv-v2-shared-link-revoke', kwargs={
                'public_id': cv.public_id,
                'link_public_id': created.data['link']['public_id'],
            }),
        )
        self.assertEqual(revoke.status_code, 204)
        self.assertEqual(self.client.get(public_url).status_code, 404)
        self.assertEqual(
            self.client.get(reverse('cv-v2-shared-link-public', kwargs={'token': 'not-a-real-token'})).status_code,
            404,
        )

        expiring = self.client.post(self.shared_links_url(cv), format='json')
        link = CvSharedLink.objects.get(public_id=expiring.data['link']['public_id'])
        CvSharedLink.objects.filter(pk=link.pk).update(expires_at=timezone.now() - timedelta(seconds=1))
        self.assertEqual(
            self.client.get(reverse('cv-v2-shared-link-public', kwargs={'token': expiring.data['token']})).status_code,
            404,
        )

    def test_owner_view_and_shared_link_management_block_idor(self):
        cv = self.create_cv()
        created = self.client.post(self.shared_links_url(cv), format='json')
        self.client.force_authenticate(self.other_candidate)

        self.assertEqual(self.client.get(self.owner_view_url(cv)).status_code, 404)
        self.assertEqual(self.client.get(self.shared_links_url(cv)).status_code, 404)
        self.assertEqual(self.client.post(self.shared_links_url(cv), format='json').status_code, 404)
        self.assertEqual(
            self.client.delete(reverse('cv-v2-shared-link-revoke', kwargs={
                'public_id': cv.public_id,
                'link_public_id': created.data['link']['public_id'],
            })).status_code,
            404,
        )

    def test_candidate_ownership_and_employer_access_are_enforced(self):
        cv = self.create_cv()
        self.client.force_authenticate(self.other_candidate)
        self.assertEqual(self.client.get(self.draft_url(cv)).status_code, 404)
        self.assertEqual(
            self.client.get(reverse('cv-v2-version-list', kwargs={'public_id': cv.public_id})).status_code,
            404,
        )

        employer = get_user_model().objects.create_user(
            email='v2-employer@example.com', password='password', role='employer',
        )
        self.client.force_authenticate(employer)
        self.assertEqual(self.client.get(self.draft_url(cv)).status_code, 403)
        self.assertEqual(self.client.get(self.owner_view_url(cv)).status_code, 403)
        self.assertEqual(self.client.get(self.shared_links_url(cv)).status_code, 403)

    def test_saved_versions_are_read_only_and_save_rolls_back_if_version_insert_fails(self):
        cv = self.create_cv()
        draft = self.client.get(self.draft_url(cv)).data
        response = self.client.patch(
            reverse('cv-v2-version-detail', kwargs={
                'public_id': cv.public_id,
                'version_public_id': cv.latest_version.public_id,
            }),
            {'plain_text': 'attempted mutation'}, format='json',
        )
        self.assertEqual(response.status_code, 405)

        with patch('apps.cvs.services.versions.CvVersion.save', side_effect=RuntimeError('database failure')):
            with self.assertRaisesRegex(RuntimeError, 'database failure'):
                self.client.post(
                    reverse('cv-v2-save-version', kwargs={'public_id': cv.public_id}),
                    format='json', HTTP_IF_MATCH=f'"lock-version-{draft["lock_version"]}"',
                )
        cv.refresh_from_db()
        self.assertEqual(CvVersion.objects.filter(cv=cv).count(), 1)
        self.assertEqual(cv.latest_version.version_kind, CvVersion.VersionKind.INITIAL)

    def test_creation_requires_verified_email_and_a_published_template_version(self):
        unverified = get_user_model().objects.create_user(
            email='v2-unverified@example.com', password='password', role='candidate', email_verified=False,
        )
        self.client.force_authenticate(unverified)
        response = self.client.post(
            reverse('cv-v2-list-create'),
            {'title': 'Blocked', 'template_public_id': self.template.public_id}, format='json',
        )
        self.assertEqual(response.status_code, 403)
        self.client.force_authenticate(self.candidate)
        self.template.current_published_version = None
        self.template.save(update_fields=['current_published_version'])
        response = self.client.post(
            reverse('cv-v2-list-create'),
            {'title': 'No release', 'template_public_id': self.template.public_id}, format='json',
        )
        self.assertEqual(response.status_code, 403)

    def test_creation_from_sample_pins_the_current_published_template_version(self):
        sample_content = empty_content('vi-VN')
        sample_content['personal_info']['full_name'] = 'Sample Candidate'
        sample = CvSampleContent.objects.create(
            locale='vi-VN',
            title='Sample content',
            content_json=sample_content,
            status=CvSampleContent.Status.PUBLISHED,
        )
        next_version = CvTemplateVersion.objects.create(
            template=self.template,
            version_number=2,
            version_status=CvTemplateVersion.VersionStatus.PUBLISHED,
            renderer_key='classic_single_column_v1',
            renderer_version='1',
            default_layout_json=empty_layout(),
            default_style_json=empty_style(),
        )
        self.template.current_published_version = next_version
        self.template.save(update_fields=['current_published_version'])

        response = self.client.post(
            reverse('cv-v2-list-create'),
            {
                'title': 'Sample CV',
                'template_public_id': self.template.public_id,
                'language': 'vi-VN',
                'sample_content_public_id': sample.public_id,
            },
            format='json',
        )

        self.assertEqual(response.status_code, 201, response.data)
        cv = UserCv.objects.get(public_id=response.data['public_id'])
        self.assertEqual(cv.current_template_version_id, next_version.id)
        self.assertEqual(cv.draft.content_json['personal_info']['full_name'], 'Sample Candidate')
        self.assertEqual(cv.latest_version.template_version_id, next_version.id)

    def test_create_blank_builds_usercv_initial_version_and_draft(self):
        response = self.client.post(
            reverse('cv-v2-list-create'),
            {'title': 'Blank CV', 'template_public_id': self.template.public_id, 'language': 'vi-VN'},
            format='json',
        )
        self.assertEqual(response.status_code, 201, response.data)
        cv = UserCv.objects.get(public_id=response.data['public_id'])
        self.assertEqual(cv.user_id, self.candidate.id)
        # Exactly one immutable initial version.
        versions = CvVersion.objects.filter(cv=cv)
        self.assertEqual(versions.count(), 1)
        self.assertEqual(cv.latest_version.version_kind, CvVersion.VersionKind.INITIAL)
        self.assertEqual(cv.latest_version.template_version_id, self.template_version.id)
        # A mutable draft seeded from the initial document.
        self.assertTrue(CvDraft.objects.filter(cv=cv).exists())
        self.assertEqual(cv.draft.content_json, cv.latest_version.content_json)

    def test_create_from_sample_builds_usercv_initial_version_and_draft(self):
        sample_content = empty_content('vi-VN')
        sample_content['personal_info']['full_name'] = 'Sample Person'
        sample = CvSampleContent.objects.create(
            locale='vi-VN', title='Sample', content_json=sample_content,
            status=CvSampleContent.Status.PUBLISHED,
        )
        response = self.client.post(
            reverse('cv-v2-list-create'),
            {
                'title': 'Sample CV', 'template_public_id': self.template.public_id,
                'language': 'vi-VN', 'sample_content_public_id': sample.public_id,
            },
            format='json',
        )
        self.assertEqual(response.status_code, 201, response.data)
        cv = UserCv.objects.get(public_id=response.data['public_id'])
        self.assertEqual(CvVersion.objects.filter(cv=cv).count(), 1)
        self.assertEqual(cv.latest_version.version_kind, CvVersion.VersionKind.INITIAL)
        self.assertTrue(CvDraft.objects.filter(cv=cv).exists())
        # Sample content flows into both the immutable baseline and the draft.
        self.assertEqual(cv.draft.content_json['personal_info']['full_name'], 'Sample Person')
        self.assertEqual(cv.latest_version.content_json['personal_info']['full_name'], 'Sample Person')

    def test_apply_sample_uses_cas_and_preserves_personal_info_and_style(self):
        cv = self.create_cv()
        cv.draft.content_json['personal_info']['full_name'] = 'Tên ứng viên thật'
        cv.draft.style_json['theme_color'] = '#2255AA'
        cv.draft.save(update_fields=['content_json', 'style_json'])
        sample_content = empty_content('vi-VN')
        sample_content['personal_info']['full_name'] = 'Tên demo không được ghi đè'
        sample_content['sections'] = [{
            'instance_id': 'sample_summary_1', 'section_key': 'summary',
            'title': 'Giới thiệu', 'enabled': True,
            'items': [{'item_id': 'sample_summary_item_1', 'value': 'Nội dung mẫu'}],
        }]
        sample = CvSampleContent.objects.create(
            locale='vi-VN', title='Sample apply', content_json=sample_content,
            status=CvSampleContent.Status.PUBLISHED,
        )
        url = reverse('cv-v2-apply-sample', kwargs={'public_id': cv.public_id})

        applied = self.client.post(
            url, {'sample_content_public_id': sample.public_id}, format='json',
            HTTP_IF_MATCH='"lock-version-0"',
        )

        self.assertEqual(applied.status_code, 200, applied.data)
        self.assertEqual(applied.data['lock_version'], 1)
        self.assertEqual(applied.data['content_json']['personal_info']['full_name'], 'Tên ứng viên thật')
        self.assertEqual(applied.data['content_json']['sections'][0]['items'][0]['value'], 'Nội dung mẫu')
        self.assertEqual(applied.data['style_json']['theme_color'], '#2255AA')
        stale = self.client.post(
            url, {'sample_content_public_id': sample.public_id}, format='json',
            HTTP_IF_MATCH='"lock-version-0"',
        )
        self.assertEqual(stale.status_code, 409)

    def test_avatar_upload_reencodes_image_and_enforces_owner_access(self):
        buffer = BytesIO()
        Image.new('RGB', (900, 700), '#2255AA').save(buffer, format='JPEG')
        upload = SimpleUploadedFile('avatar.jpg', buffer.getvalue(), content_type='image/jpeg')

        created = self.client.post(reverse('cv-v2-asset-upload'), {'file': upload}, format='multipart')

        self.assertEqual(created.status_code, 201, created.data)
        self.assertNotIn('storage_key', created.data)
        self.assertLessEqual(created.data['width'], 512)
        self.assertLessEqual(created.data['height'], 512)
        self.assertIn('?token=', created.data['url'])
        asset = CvAsset.objects.get(public_id=created.data['public_id'])
        content_url = reverse('cv-v2-asset-content', kwargs={'asset_public_id': asset.public_id})
        self.client.force_authenticate(self.other_candidate)
        self.assertEqual(self.client.get(content_url).status_code, 404)
        self.assertEqual(self.client.get(created.data['url']).status_code, 200)
        self.client.force_authenticate(self.candidate)
        self.assertEqual(self.client.get(content_url).status_code, 200)

        invalid = SimpleUploadedFile('avatar.png', b'not-an-image', content_type='image/png')
        rejected = self.client.post(reverse('cv-v2-asset-upload'), {'file': invalid}, format='multipart')
        self.assertEqual(rejected.status_code, 400)

    def test_create_from_position_resolves_blueprint_and_persists_taxonomy_identity(self):
        position = JobCategory.objects.create(name='Nhân viên CSKH')
        JobCategoryLocalization.objects.bulk_create([
            JobCategoryLocalization(category=position, locale='vi-VN', display_name='Nhân viên CSKH'),
            JobCategoryLocalization(category=position, locale='en-US', display_name='Customer Service Representative'),
        ])
        blueprint = CvContentBlueprint.objects.get(locale='en-US', experience_level='unspecified')
        blueprint.summary_template = 'Build a career as {position}.'
        blueprint.save(update_fields=['summary_template', 'updated_at'])

        response = self.client.post(
            reverse('cv-v2-list-create'),
            {
                'title': 'Customer Service CV',
                'template_public_id': self.template.public_id,
                'language': 'en-US',
                'position_public_id': position.public_id,
            },
            format='json',
        )

        self.assertEqual(response.status_code, 201, response.data)
        cv = UserCv.objects.get(public_id=response.data['public_id'])
        self.assertEqual(cv.position_id, position.id)
        self.assertEqual(response.data['position_public_id'], position.public_id)
        self.assertEqual(cv.draft.content_json['personal_info']['headline'], 'Customer Service Representative')

    def test_legacy_dual_write_does_not_discard_a_v2_layout(self):
        cv = self.create_cv()
        draft = self.client.get(self.draft_url(cv)).data
        payload = {
            key: draft[key]
            for key in ('schema_version', 'content_json', 'layout_json', 'style_json')
        }
        payload['layout_json']['page']['margin_mm'] = 16
        self.client.put(self.draft_url(cv), payload, format='json', HTTP_IF_MATCH='"lock-version-0"')

        cv.refresh_from_db()
        cv.title = 'Changed through legacy metadata flow'
        cv.save(update_fields=['title'])
        sync_legacy_builder_draft(cv, self.candidate)

        self.assertEqual(cv.draft.layout_json['page']['margin_mm'], 16)

    def test_pdf_export_uses_a_selected_immutable_version_not_the_mutable_draft(self):
        cv = self.create_cv()
        draft = self.client.get(self.draft_url(cv)).data
        saved_payload = {key: deepcopy(draft[key]) for key in ('schema_version', 'content_json', 'layout_json', 'style_json')}
        saved_payload['content_json']['personal_info']['full_name'] = 'Frozen Export Candidate'
        self.client.put(self.draft_url(cv), saved_payload, format='json', HTTP_IF_MATCH='"lock-version-0"')
        saved = self.client.post(
            reverse('cv-v2-save-version', kwargs={'public_id': cv.public_id}),
            format='json', HTTP_IF_MATCH='"lock-version-1"',
        )
        self.assertEqual(saved.status_code, 201, saved.data)

        mutable_payload = deepcopy(saved_payload)
        mutable_payload['content_json']['personal_info']['full_name'] = 'Draft Must Never Export'
        self.client.put(self.draft_url(cv), mutable_payload, format='json', HTTP_IF_MATCH='"lock-version-1"')

        with patch('apps.cvs.services.exports._dispatch_after_commit'):
            response = self.client.post(
                self.exports_url(cv), {'version_public_id': saved.data['public_id']}, format='json',
                REMOTE_ADDR='203.0.113.10', HTTP_USER_AGENT='export-browser',
            )
        self.assertEqual(response.status_code, 201, response.data)
        export = CvExport.objects.get(public_id=response.data['public_id'])
        self.assertEqual(export.status, CvExport.Status.PENDING)
        self.assertEqual(export.version.public_id, saved.data['public_id'])
        self.assertEqual(export.version.content_json['personal_info']['full_name'], 'Frozen Export Candidate')
        self.assertEqual(cv.draft.content_json['personal_info']['full_name'], 'Draft Must Never Export')
        self.assertEqual(export.renderer_key, 'classic_single_column_v1')
        self.assertEqual(export.renderer_version, '1')
        self.assertNotIn('content_json', export.render_config)
        self.assertNotIn('storage_key', response.data)
        self.assertIsNone(response.data['download_url'])

        html = build_cv_pdf_html(export.version)
        self.assertIn('Frozen Export Candidate', html)
        self.assertNotIn('Draft Must Never Export', html)
        self.assertIn('@page { size: A4;', html)
        self.assertIn('break-inside: avoid-page', html)
        audit = CvAccessLog.objects.get(cv=cv, access_channel=CvAccessLog.AccessChannel.EXPORT)
        self.assertEqual(audit.version_id, export.version_id)
        self.assertEqual(audit.actor_user_id, self.candidate.id)
        self.assertEqual(len(audit.ip_hash), 64)
        self.assertNotIn('content_json', {field.name for field in CvAccessLog._meta.fields})

    def test_pdf_export_defaults_to_published_before_newer_latest_version(self):
        cv = self.create_cv()
        published = self.client.post(
            reverse('cv-v2-publish', kwargs={'public_id': cv.public_id}),
            format='json', HTTP_IF_MATCH='"lock-version-0"',
        )
        self.assertEqual(published.status_code, 201, published.data)
        saved = self.client.post(
            reverse('cv-v2-save-version', kwargs={'public_id': cv.public_id}),
            format='json', HTTP_IF_MATCH='"lock-version-0"',
        )
        self.assertEqual(saved.status_code, 201, saved.data)
        self.assertNotEqual(saved.data['public_id'], published.data['public_id'])

        with patch('apps.cvs.services.exports._dispatch_after_commit'):
            response = self.client.post(self.exports_url(cv), format='json')
        self.assertEqual(response.status_code, 201, response.data)
        self.assertEqual(response.data['version_public_id'], published.data['public_id'])

    def test_pdf_renderer_uses_the_pinned_two_column_contract_without_content_conversion(self):
        cv = self.create_cv()
        alternate_template, alternate_version = self.create_alternate_template()
        switched = self.client.put(
            reverse('cv-v2-template-switch', kwargs={'public_id': cv.public_id}),
            {'template_public_id': alternate_template.public_id},
            format='json', HTTP_IF_MATCH='"lock-version-0"',
        )
        self.assertEqual(switched.status_code, 200, switched.data)
        draft = self.client.get(self.draft_url(cv)).data
        payload = {
            key: deepcopy(draft[key])
            for key in ('schema_version', 'content_json', 'layout_json', 'style_json')
        }
        payload['content_json']['sections'] = [{
            'instance_id': 'summary_1',
            'section_key': 'summary',
            'title': 'Mục tiêu nghề nghiệp',
            'enabled': True,
            'items': [{'item_id': 'summary_item_1', 'value': 'Xây dựng sản phẩm hữu ích.'}],
        }]
        payload['layout_json']['regions'][0]['section_instance_ids'] = ['summary_1']
        updated = self.client.put(
            self.draft_url(cv), payload, format='json', HTTP_IF_MATCH='"lock-version-1"',
        )
        self.assertEqual(updated.status_code, 200, updated.data)
        saved = self.client.post(
            reverse('cv-v2-save-version', kwargs={'public_id': cv.public_id}),
            format='json', HTTP_IF_MATCH='"lock-version-2"',
        )
        version = CvVersion.objects.get(public_id=saved.data['public_id'])

        html = build_cv_pdf_html(version)

        self.assertEqual(version.template_version_id, alternate_version.id)
        self.assertIn('data-renderer="classic_two_column_v1"', html)
        self.assertIn('data-renderer-version="1"', html)
        self.assertIn('class="region region-main" style="width: 68%;"', html)
        self.assertIn('class="region region-sidebar" style="width: 32%;"', html)
        self.assertIn('.region { float: left;', html)
        self.assertEqual(version.content_json, cv.draft.content_json)

        pdf_text = ''.join(
            page.extract_text() or ''
            for page in PdfReader(BytesIO(render_cv_version_pdf(version))).pages
        )
        self.assertIn('MỤC TIÊU NGHỀ NGHIỆP', pdf_text.upper())

    def test_pdf_renderer_matches_avatar_size_and_omits_empty_sections(self):
        cv = self.create_cv()
        draft = self.client.get(self.draft_url(cv)).data
        payload = {key: deepcopy(draft[key]) for key in ('schema_version', 'content_json', 'layout_json', 'style_json')}
        payload['content_json']['personal_info']['avatar_size_mm'] = 36
        payload['content_json']['personal_info']['avatar_zoom'] = 1.6
        payload['content_json']['sections'].append({
            'instance_id': 'empty_summary_1',
            'section_key': 'summary',
            'title': 'EMPTY SECTION MUST BE HIDDEN',
            'enabled': True,
            'items': [{'item_id': 'empty_summary_item_1', 'value': ''}],
        })
        payload['layout_json']['regions'][0]['section_instance_ids'].append('empty_summary_1')
        updated = self.client.put(
            self.draft_url(cv), payload, format='json', HTTP_IF_MATCH='"lock-version-0"',
        )
        self.assertEqual(updated.status_code, 200, updated.data)
        saved = self.client.post(
            reverse('cv-v2-save-version', kwargs={'public_id': cv.public_id}),
            format='json', HTTP_IF_MATCH='"lock-version-1"',
        )
        version = CvVersion.objects.get(public_id=saved.data['public_id'])

        html = build_cv_pdf_html(version)

        self.assertIn('height: 36mm', html)
        self.assertIn('width: 36mm', html)
        self.assertIn('transform: scale(1.6)', html)
        self.assertNotIn('EMPTY SECTION MUST BE HIDDEN', html)

    def test_pdf_export_reuses_valid_artifact_and_download_url_is_owner_controlled(self):
        cv = self.create_cv()
        with patch('apps.cvs.services.exports._dispatch_after_commit'):
            created = self.client.post(self.exports_url(cv), format='json')
        self.assertEqual(created.status_code, 201, created.data)
        export = CvExport.objects.get(public_id=created.data['public_id'])
        key = f'private-test/{export.public_id}.pdf'
        default_storage.save(key, ContentFile(b'%PDF-1.4 immutable artifact'))
        CvExport.objects.filter(pk=export.pk).update(
            status=CvExport.Status.COMPLETED,
            storage_key=key,
            file_size_bytes=27,
            checksum_sha256=sha256(b'%PDF-1.4 immutable artifact').hexdigest(),
            completed_at=timezone.now(),
        )
        export.refresh_from_db()

        reused = self.client.post(self.exports_url(cv), format='json')
        self.assertEqual(reused.status_code, 200, reused.data)
        self.assertEqual(reused.data['public_id'], export.public_id)
        self.assertEqual(CvExport.objects.filter(cv=cv).count(), 1)
        self.assertTrue(reused.data['download_url'].endswith(self.export_url(cv, export, 'download')))
        self.assertNotIn(key, reused.data['download_url'])
        self.assertNotIn('storage_key', reused.data)

        download = self.client.get(self.export_url(cv, export, 'download'))
        self.assertEqual(download.status_code, 200)
        self.assertEqual(download['Content-Type'], 'application/pdf')
        self.assertEqual(b''.join(download.streaming_content), b'%PDF-1.4 immutable artifact')

    def test_pdf_export_worker_marks_failure_and_retry_keeps_the_same_version(self):
        cv = self.create_cv()
        with patch('apps.cvs.services.exports._dispatch_after_commit'):
            created = self.client.post(self.exports_url(cv), format='json')
        export = CvExport.objects.get(public_id=created.data['public_id'])
        version_id = export.version_id
        draft_before = deepcopy(cv.draft.content_json)

        with patch('apps.cvs.tasks.render_cv_version_pdf', side_effect=RuntimeError('renderer unavailable')):
            render_cv_export_job(export.pk)
        export.refresh_from_db()
        self.assertEqual(export.status, CvExport.Status.FAILED)
        self.assertEqual(export.last_error, 'render_failed')
        self.assertEqual(export.version_id, version_id)
        cv.refresh_from_db()
        self.assertEqual(cv.draft.content_json, draft_before)
        self.assertEqual(CvVersion.objects.filter(cv=cv).count(), 1)

        with patch('apps.cvs.services.exports._dispatch_after_commit'):
            retry = self.client.post(self.export_url(cv, export, 'retry'), format='json')
        self.assertEqual(retry.status_code, 202, retry.data)
        export.refresh_from_db()
        self.assertEqual(export.status, CvExport.Status.PENDING)
        self.assertEqual(export.version_id, version_id)
        self.assertEqual(CvExport.objects.filter(cv=cv).count(), 1)

    def test_pdf_export_worker_stores_a_private_pdf_artifact_from_version_only(self):
        cv = self.create_cv()
        with patch('apps.cvs.services.exports._dispatch_after_commit'):
            created = self.client.post(self.exports_url(cv), format='json')
        export = CvExport.objects.get(public_id=created.data['public_id'])
        with patch('apps.cvs.tasks.render_cv_version_pdf', return_value=b'%PDF-1.4 rendered version') as renderer:
            render_cv_export_job(export.pk)
        export.refresh_from_db()
        self.assertEqual(export.status, CvExport.Status.COMPLETED)
        self.assertTrue(export.storage_key.startswith(f'cvs/exports/{cv.public_id}/{export.version.public_id}/'))
        self.assertTrue(default_storage.exists(export.storage_key))
        self.assertEqual(renderer.call_args.args[0].pk, export.version_id)
        self.assertNotIn('draft', renderer.call_args.kwargs)

    def test_pdf_export_blocks_idor_and_foreign_version_selection(self):
        cv = self.create_cv()
        with patch('apps.cvs.services.exports._dispatch_after_commit'):
            created = self.client.post(self.exports_url(cv), format='json')
        export = CvExport.objects.get(public_id=created.data['public_id'])

        self.client.force_authenticate(self.other_candidate)
        self.assertEqual(self.client.get(self.exports_url(cv)).status_code, 404)
        self.assertEqual(self.client.post(self.exports_url(cv), format='json').status_code, 404)
        self.assertEqual(self.client.get(self.export_url(cv, export)).status_code, 404)
        self.assertEqual(self.client.get(self.export_url(cv, export, 'download')).status_code, 404)

        self.client.force_authenticate(self.candidate)
        other_cv = UserCv.objects.create(
            user=self.other_candidate,
            title='Other CV',
            cv_type=UserCv.CvType.BUILDER,
            source=UserCv.Source.BUILDER,
            language='vi-VN',
        )
        other_version = CvVersion.objects.create(
            cv=other_cv,
            version_number=1,
            version_kind=CvVersion.VersionKind.INITIAL,
            template_version=self.template_version,
            schema_version=1,
            content_json=empty_content(),
            layout_json=empty_layout(),
            style_json=empty_style(),
            content_hash='a' * 64,
            created_by=self.other_candidate,
        )
        response = self.client.post(
            self.exports_url(cv), {'version_public_id': other_version.public_id}, format='json',
        )
        self.assertEqual(response.status_code, 400)
        self.assertEqual(CvExport.objects.filter(cv=cv).count(), 1)

        employer = get_user_model().objects.create_user(
            email='export-employer@example.com', password='password', role='employer',
        )
        self.client.force_authenticate(employer)
        self.assertEqual(self.client.post(self.exports_url(cv), format='json').status_code, 403)

    @patch('apps.cvs.tasks.first_pdf_page_image', return_value=b'webp-preview')
    @patch('apps.cvs.tasks.render_cv_version_pdf', return_value=b'%PDF-1.4')
    def test_private_thumbnail_is_generated_from_latest_version_and_owner_scoped(self, _render, _raster):
        cv = self.create_cv()
        cv.refresh_from_db()

        generate_cv_thumbnail(cv.latest_version_id)
        cv.refresh_from_db()

        self.assertTrue(cv.thumbnail_url.endswith('.webp'))
        response = self.client.get(reverse('cv-v2-thumbnail', kwargs={'public_id': cv.public_id}))
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response['Content-Type'], 'image/webp')
        self.assertEqual(b''.join(response.streaming_content), b'webp-preview')

        self.client.force_authenticate(self.other_candidate)
        self.assertEqual(
            self.client.get(reverse('cv-v2-thumbnail', kwargs={'public_id': cv.public_id})).status_code,
            404,
        )
