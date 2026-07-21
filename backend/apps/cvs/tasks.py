"""Celery tasks for immutable CV export artifacts."""

from __future__ import annotations

import logging
from hashlib import sha256
from time import monotonic

from celery import shared_task
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from django.db import transaction
from django.utils import timezone

from apps.ai_core.services import AiCvParseError, structure_cv_text
from common.metrics import record_metric
from common.pdf_raster import first_pdf_page_image

from .models import CvExport, CvImportJob, UserCv
from .services.composition import compose_cv_document, overlay_actor_identity
from .services.pdf_renderer import render_cv_version_pdf
from .services.thumbnails import thumbnail_key_for
from .services.versions import create_version

logger = logging.getLogger(__name__)


@shared_task
def generate_cv_thumbnail(version_id: int) -> None:
    """Create a private WebP preview from the newest immutable saved version."""
    from .models import CvVersion

    started_at = monotonic()
    version = (
        CvVersion.objects.select_related('cv', 'template_version').filter(pk=version_id).first()
    )
    if version is None or version.cv_id is None or version.cv.latest_version_id != version.pk:
        return
    cv = version.cv
    storage_key = thumbnail_key_for(cv, version)
    try:
        if default_storage.exists(storage_key):
            saved_key = storage_key
        else:
            pdf_bytes = render_cv_version_pdf(version)
            image_bytes = first_pdf_page_image(
                pdf_bytes, width=900, image_format='WEBP', quality=88
            )
            saved_key = default_storage.save(storage_key, ContentFile(image_bytes))
    except Exception:  # noqa: BLE001 - never log CV content or renderer payload
        logger.warning('Private CV thumbnail generation failed for version %s.', version_id)
        record_metric('cv_snapshot_failure')
        record_metric(
            'cv_snapshot_duration_ms', round((monotonic() - started_at) * 1000, 2), status='failed'
        )
        return

    with transaction.atomic():
        current = UserCv.objects.select_for_update().filter(pk=cv.pk).first()
        if current is None or current.latest_version_id != version.pk:
            default_storage.delete(saved_key)
            return
        old_key = current.thumbnail_url
        current.thumbnail_url = saved_key
        current.save(update_fields=['thumbnail_url', 'updated_at'])
        if old_key and old_key != saved_key:
            transaction.on_commit(lambda: default_storage.delete(old_key))
    record_metric(
        'cv_snapshot_duration_ms', round((monotonic() - started_at) * 1000, 2), status='completed'
    )


def _storage_key(export: CvExport) -> str:
    # Attempt number prevents a retry from silently serving a stale partial
    # artifact; the key remains private and is never serialized to clients.
    return (
        f'cvs/exports/{export.cv.public_id}/{export.version.public_id}/'
        f'{export.public_id}/attempt-{export.attempts}.pdf'
    )


@shared_task
def render_cv_export_job(export_id: int) -> None:
    """Render a queued immutable version and atomically record its private artifact."""
    with transaction.atomic():
        export = (
            # ``version.template_version`` is nullable for legacy rows. Lock
            # the job row only; PostgreSQL rejects FOR UPDATE on that nullable
            # outer-join side.
            CvExport.objects.select_for_update(of=('self',))
            .select_related('cv', 'version__template_version')
            .filter(pk=export_id)
            .first()
        )
        if export is None or export.status != CvExport.Status.PENDING:
            return
        export.status = CvExport.Status.PROCESSING
        export.attempts += 1
        export.started_at = timezone.now()
        export.save(update_fields=['status', 'attempts', 'started_at', 'updated_at'])

    try:
        # The renderer receives only the frozen version relation. Drafts and
        # frontend DOM are intentionally absent from this worker path.
        pdf_bytes = render_cv_version_pdf(export.version)
        storage_key = _storage_key(export)
        saved_key = default_storage.save(storage_key, ContentFile(pdf_bytes))
        checksum = sha256(pdf_bytes).hexdigest()
    except Exception:  # noqa: BLE001 - preserve no renderer exception payload that could contain PII
        logger.warning('Immutable CV PDF export %s failed.', export_id)
        CvExport.objects.filter(pk=export_id, status=CvExport.Status.PROCESSING).update(
            status=CvExport.Status.FAILED,
            last_error='render_failed',
            failed_at=timezone.now(),
        )
        return

    now = timezone.now()
    with transaction.atomic():
        completed = (
            CvExport.objects.select_for_update()
            .filter(
                pk=export_id,
                status=CvExport.Status.PROCESSING,
            )
            .first()
        )
        if completed is None:
            # A manual repair/retry changed state while the worker rendered.
            # Retain no public reference to this unclaimed storage object.
            if default_storage.exists(saved_key):
                default_storage.delete(saved_key)
            return
        completed.status = CvExport.Status.COMPLETED
        completed.storage_key = saved_key
        completed.file_size_bytes = len(pdf_bytes)
        completed.checksum_sha256 = checksum
        completed.last_error = ''
        completed.completed_at = now
        completed.failed_at = None
        completed.save(
            update_fields=[
                'status',
                'storage_key',
                'file_size_bytes',
                'checksum_sha256',
                'last_error',
                'completed_at',
                'failed_at',
                'updated_at',
            ]
        )
        UserCv.objects.filter(pk=completed.cv_id).update(last_exported_at=now, updated_at=now)


@shared_task
def dispatch_pending_cv_export_jobs() -> None:
    """Recover jobs left pending when a web process could not reach the broker."""
    job_ids = list(
        CvExport.objects.filter(status=CvExport.Status.PENDING)
        .order_by('queued_at')
        .values_list('pk', flat=True)[:100]
    )
    for export_id in job_ids:
        try:
            render_cv_export_job.delay(export_id)
        except Exception:  # noqa: BLE001 - leave the durable row for the next sweep
            logger.warning('Unable to dispatch pending immutable CV PDF export %s.', export_id)


class ImportProcessingError(ValueError):
    def __init__(self, code):
        self.code = code
        super().__init__(code)


def _extract_import_text(cv):
    try:
        with default_storage.open(cv.file_url, 'rb') as source:
            if cv.file_type == 'pdf':
                from pypdf import PdfReader

                reader = PdfReader(source)
                if len(reader.pages) > 20:
                    raise ImportProcessingError('too_many_pages')
                text = '\n'.join(page.extract_text() or '' for page in reader.pages)
            elif cv.file_type == 'docx':
                from docx import Document

                document = Document(source)
                paragraphs = [paragraph.text for paragraph in document.paragraphs]
                table_cells = [
                    cell.text
                    for table in document.tables
                    for row in table.rows
                    for cell in row.cells
                ]
                text = '\n'.join(paragraphs + table_cells)
            else:
                raise ImportProcessingError('unsupported_file_type')
    except ImportProcessingError:
        raise
    except Exception as error:
        raise ImportProcessingError('text_extraction_failed') from error
    text = text.strip()
    if cv.file_type == 'pdf' and len(text) < 40:
        raise ImportProcessingError('scanned_pdf_ocr_unavailable')
    if len(text) < 10:
        raise ImportProcessingError('insufficient_text')
    if len(text) > 100_000:
        raise ImportProcessingError('text_too_long')
    return text


def _fail_import(job_id, code):
    now = timezone.now()
    with transaction.atomic():
        job = CvImportJob.objects.select_for_update().select_related('cv').filter(pk=job_id).first()
        if job is None or job.status != CvImportJob.Status.PROCESSING:
            return
        job.status = CvImportJob.Status.FAILED
        job.failure_code = code
        job.failed_at = now
        job.save(update_fields=['status', 'failure_code', 'failed_at', 'updated_at'])
        UserCv.objects.filter(pk=job.cv_id).update(
            status=UserCv.Status.FAILED,
            processing_status=UserCv.ProcessingStatus.FAILED,
            error_message=code,
            updated_at=now,
        )
        duration_ms = (now - job.started_at).total_seconds() * 1000 if job.started_at else 0
        record_metric('cv_import_failure', failure_code=code)
        record_metric('cv_import_duration_ms', round(duration_ms, 2), status='failed')


@shared_task(soft_time_limit=50, time_limit=60)
def process_cv_import_job(job_id):
    with transaction.atomic():
        job = (
            CvImportJob.objects.select_for_update(of=('self',))
            .select_related(
                'cv__template__current_published_version',
                'user',
            )
            .filter(pk=job_id)
            .first()
        )
        if job is None or job.status != CvImportJob.Status.QUEUED:
            return
        job.status = CvImportJob.Status.PROCESSING
        job.attempts += 1
        job.started_at = timezone.now()
        job.save(update_fields=['status', 'attempts', 'started_at', 'updated_at'])
        UserCv.objects.filter(pk=job.cv_id).update(
            status=UserCv.Status.PROCESSING,
            processing_status=UserCv.ProcessingStatus.PROCESSING,
            updated_at=timezone.now(),
        )

    try:
        text_value = _extract_import_text(job.cv)
        content = structure_cv_text(text_value, job.cv.language)
        content = overlay_actor_identity(content, job.user, fill_only=True)
        theme_color = (job.cv.style_config or {}).get('theme_color')
        document = compose_cv_document(
            template=job.cv.template,
            content_json=content,
            theme_color=theme_color,
        )
        with transaction.atomic():
            current = (
                CvImportJob.objects.select_for_update().select_related('cv', 'user').get(pk=job_id)
            )
            if current.status != CvImportJob.Status.PROCESSING:
                return
            version = create_version(
                cv=current.cv,
                actor=current.user,
                content_json=document['content_json'],
                layout_json=document['layout_json'],
                style_json=document['style_json'],
                version_kind='imported',
                template_version=current.cv.template.current_published_version,
                create_or_replace_draft=True,
            )
            now = timezone.now()
            UserCv.objects.filter(pk=current.cv_id).update(
                cv_data=document['content_json'],
                style_config=document['style_json'],
                status=UserCv.Status.ANALYZED,
                processing_status=UserCv.ProcessingStatus.ANALYZED,
                error_message='',
                raw_text='',
                normalized_text='',
                last_analyzed_at=now,
                updated_at=now,
            )
            current.status = CvImportJob.Status.COMPLETED
            current.result_version = version
            current.failure_code = ''
            current.completed_at = now
            current.failed_at = None
            current.save(
                update_fields=[
                    'status',
                    'result_version',
                    'failure_code',
                    'completed_at',
                    'failed_at',
                    'updated_at',
                ]
            )
            duration_ms = (
                (now - current.started_at).total_seconds() * 1000 if current.started_at else 0
            )
            record_metric('cv_import_duration_ms', round(duration_ms, 2), status='completed')
    except ImportProcessingError as error:
        _fail_import(job_id, error.code)
    except AiCvParseError as error:
        _fail_import(job_id, error.code)
    except Exception:  # noqa: BLE001 - never log raw text or provider response
        logger.warning('CV import job %s failed.', job_id)
        _fail_import(job_id, 'import_processing_failed')


@shared_task
def dispatch_pending_cv_import_jobs():
    for job_id in CvImportJob.objects.filter(status=CvImportJob.Status.QUEUED).values_list(
        'pk', flat=True
    )[:100]:
        process_cv_import_job.delay(job_id)


@shared_task
def purge_expired_cv_import_sources():
    for job in (
        CvImportJob.objects.select_related('cv')
        .filter(
            source_expires_at__lte=timezone.now(),
            status__in=[CvImportJob.Status.COMPLETED, CvImportJob.Status.FAILED],
        )
        .exclude(cv__file_url='')
        .iterator()
    ):
        storage_key = job.cv.file_url
        if default_storage.exists(storage_key):
            default_storage.delete(storage_key)
        UserCv.objects.filter(pk=job.cv_id, file_url=storage_key).update(
            file_url='', updated_at=timezone.now()
        )
